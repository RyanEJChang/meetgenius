import io
import re

from docx import Document
from docx.shared import Inches


def _safe_language_label(locale_code: str) -> str:
    """以傳入 locale 原樣作為顯示用標籤；僅做基本清理。"""
    return (locale_code or '').strip()


def apply_speaker_names_to_text(text: str, speaker_name_mapping: dict) -> str:
    """
    將發言者名稱映射應用到文字內容中
    
    Args:
        text (str): 原始文字內容
        speaker_name_mapping (dict): 發言者名稱映射
    
    Returns:
        str: 應用發言者名稱映射後的文字內容
    """
    if not speaker_name_mapping or not text:
        return text
    
    modified_text = text
    for original_speaker, custom_name in speaker_name_mapping.items():
        # 替換各種可能的發言者格式
        patterns = [
            rf'{re.escape(original_speaker)}：',  # 發言者00：
            rf'{re.escape(original_speaker)}:',   # 發言者00:
            rf'\b{re.escape(original_speaker)}\b',  # 單獨的發言者名稱
        ]
        
        for pattern in patterns:
            modified_text = re.sub(pattern, custom_name, modified_text)
    
    return modified_text


_OUTCOME_LABELS = {
    'consensus': '已達成共識',
    'unresolved': '尚未解決',
    'deferred': '延後處理',
    'clarified': '已釐清',
}

_OWNER_SUFFIX = {
    'external': '（外部單位）',
    'team': '（團隊）',
    'person': '',
    'unassigned': '',
}


def _outcome_label(outcome) -> str:
    return _OUTCOME_LABELS.get(outcome or '', '尚未解決')


def _owner_suffix(owner_type) -> str:
    return _OWNER_SUFFIX.get(owner_type or '', '')


def format_summary_for_export(summary_data: dict, meeting_info, speaker_name_mapping: dict = None) -> str:
    """
    將摘要資料格式化為純文字字串。
    
    Args:
        summary_data (dict): 摘要資料
        meeting_info: 會議資訊
        speaker_name_mapping (dict): 發言者名稱映射
    """
    lines = []
    # 安全地獲取會議檔名
    original_filename = meeting_info.get('original_filename', 'N/A')
    lines.append(f"會議摘要：{original_filename}")
    lines.append("=" * 40)
    # lines.append(f"生成時間：{summary_data.get('summary_generated_at', 'N/A')}")
    lines.append("\n")

    lines.append("【會議總結】")
    lines.append("-" * 20)
    summary = summary_data.get('summary', '無內容')
    if speaker_name_mapping:
        summary = apply_speaker_names_to_text(summary, speaker_name_mapping)
    lines.append(summary)
    lines.append("\n")

    decisions = summary_data.get('decisions', [])
    if decisions:
        lines.append("【決議】")
        lines.append("-" * 20)
        for d in decisions:
            if not isinstance(d, dict):
                lines.append(f"• {d}")
                continue
            text = d.get('decision', '')
            if speaker_name_mapping:
                text = apply_speaker_names_to_text(text, speaker_name_mapping)
            lines.append(f"• {text}")
            for r in d.get('ruled_out') or []:
                lines.append(f"    （已排除）{r}")
        lines.append("\n")

    positions = summary_data.get('positions', [])
    if positions:
        lines.append("【各方立場】")
        lines.append("-" * 20)
        for p in positions:
            if not isinstance(p, dict):
                continue
            lines.append(f"• {p.get('topic', '')}　[{_outcome_label(p.get('outcome'))}]")
            for s in p.get('stances') or []:
                who = s.get('speaker', '')
                if speaker_name_mapping:
                    who = apply_speaker_names_to_text(who, speaker_name_mapping)
                lines.append(f"    - {who}：{s.get('position', '')}")
        lines.append("\n")

    lines.append("【行動項目】")
    lines.append("-" * 20)
    action_items = summary_data.get('action_items', [])
    if action_items:
        for item in action_items:
            if not isinstance(item, dict):
                lines.append(f"• {item}")
                continue
            item_text = item.get('item', '')
            if speaker_name_mapping:
                item_text = apply_speaker_names_to_text(item_text, speaker_name_mapping)
            lines.append(f"• {item_text}")

            owner = item.get('owner')
            if speaker_name_mapping and owner:
                owner = apply_speaker_names_to_text(owner, speaker_name_mapping)
            owner_label = f"{owner}{_owner_suffix(item.get('owner_type'))}" if owner else "未指派"
            lines.append(f"    負責：{owner_label}")

            if item.get('precondition'):
                lines.append(f"    前提：{item['precondition']}")
    else:
        lines.append("無內容")
    lines.append("\n")

    lines.append("【關鍵字】")
    lines.append("-" * 20)
    keywords = summary_data.get('keywords', [])
    if keywords:
        # 確保關鍵字列表中的每個項目都是字串
        def to_str(k):
            if isinstance(k, dict):
                return k.get('keyword', str(k))
            return str(k)
        
        keyword_list = [to_str(k) for k in keywords]
        keyword_text = "、".join(keyword_list)

        if speaker_name_mapping:
            keyword_text = apply_speaker_names_to_text(keyword_text, speaker_name_mapping)
        lines.append(keyword_text)
    else:
        lines.append("無內容")

    return "\n".join(lines)


def create_summary_docx(summary_data: dict, meeting_info, speaker_name_mapping: dict = None) -> io.BytesIO:
    """
    從摘要資料創建 Word 文件。
    
    Args:
        summary_data (dict): 摘要資料
        meeting_info: 會議資訊
        speaker_name_mapping (dict): 發言者名稱映射
    """
    document = Document()
    
    # 安全地獲取會議檔名
    original_filename = meeting_info.get('original_filename', 'N/A')
    
    # 設置文件標題
    document.add_heading(f"會議摘要：{original_filename}", level=0)
    
    # 添加生成時間
    # p = document.add_paragraph()
    # p.add_run('生成時間：').italic = True
    # p.add_run(str(summary_data.get('summary_generated_at', 'N/A')))
    
    # 添加總結
    document.add_heading('會議總結', level=1)
    summary = summary_data.get('summary', '無內容')
    if speaker_name_mapping:
        summary = apply_speaker_names_to_text(summary, speaker_name_mapping)
    document.add_paragraph(summary)

    # 添加決議
    decisions = summary_data.get('decisions', [])
    if decisions:
        document.add_heading('決議', level=1)
        for d in decisions:
            if not isinstance(d, dict):
                document.add_paragraph(str(d), style='List Bullet')
                continue
            text = d.get('decision', '')
            if speaker_name_mapping:
                text = apply_speaker_names_to_text(text, speaker_name_mapping)
            document.add_paragraph(text, style='List Bullet')
            for r in d.get('ruled_out') or []:
                p = document.add_paragraph(style='List Bullet 2')
                run = p.add_run(f"已排除：{r}")
                run.italic = True

    # 添加各方立場
    positions = summary_data.get('positions', [])
    if positions:
        document.add_heading('各方立場', level=1)
        for pos in positions:
            if not isinstance(pos, dict):
                continue
            p = document.add_paragraph(style='List Bullet')
            p.add_run(pos.get('topic', '')).bold = True
            p.add_run(f"　[{_outcome_label(pos.get('outcome'))}]").italic = True
            for s in pos.get('stances') or []:
                who = s.get('speaker', '')
                if speaker_name_mapping:
                    who = apply_speaker_names_to_text(who, speaker_name_mapping)
                sp = document.add_paragraph(style='List Bullet 2')
                sp.add_run(f"{who}：").bold = True
                sp.add_run(s.get('position', ''))

    # 添加行動項目
    document.add_heading('行動項目', level=1)
    action_items = summary_data.get('action_items', [])
    if action_items:
        for item in action_items:
            if not isinstance(item, dict):
                document.add_paragraph(str(item), style='List Bullet')
                continue

            item_text = item.get('item', '')
            if speaker_name_mapping:
                item_text = apply_speaker_names_to_text(item_text, speaker_name_mapping)
            document.add_paragraph(item_text, style='List Bullet')

            owner = item.get('owner')
            if speaker_name_mapping and owner:
                owner = apply_speaker_names_to_text(owner, speaker_name_mapping)
            detail = document.add_paragraph(style='List Bullet 2')
            detail.add_run('負責：').bold = True
            detail.add_run(f"{owner}{_owner_suffix(item.get('owner_type'))}" if owner else '未指派')

            if item.get('precondition'):
                pre = document.add_paragraph(style='List Bullet 2')
                pre.add_run('前提：').bold = True
                pre.add_run(item['precondition'])
    else:
        document.add_paragraph('無內容')

    # 添加關鍵字
    document.add_heading('關鍵字', level=1)
    keywords = summary_data.get('keywords', [])
    if keywords:
        # 確保關鍵字列表中的每個項目都是字串
        def to_str(k):
            if isinstance(k, dict):
                return k.get('keyword', str(k))
            return str(k)
            
        keyword_list = [to_str(k) for k in keywords]
        keyword_text = "、".join(keyword_list)

        if speaker_name_mapping:
            keyword_text = apply_speaker_names_to_text(keyword_text, speaker_name_mapping)
        document.add_paragraph(keyword_text)
    else:
        document.add_paragraph('無內容')

    # 將文件保存在內存中
    file_stream = io.BytesIO()
    document.save(file_stream)
    file_stream.seek(0)
    
    return file_stream


def create_transcript_docx(
    transcript_segments: list,
    meeting_info: dict
) -> io.BytesIO:
    """
    從單語逐字稿段落建立 Word 文件。

    Args:
        transcript_segments (list): 由 parse_srt_content 輸出之段落列表
        meeting_info (dict): 會議資訊
    Returns:
        io.BytesIO: 內存中的 .docx 檔案流
    """
    document = Document()

    original_filename = meeting_info.get('original_filename', 'N/A')
    document.add_heading(f"逐字稿：{original_filename}", level=0)

    table = document.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '時間'
    hdr_cells[1].text = '發言者'
    hdr_cells[2].text = '內容'

    for seg in transcript_segments:
        row = table.add_row().cells
        row[0].text = seg.get('start_time_formatted', '')
        row[1].text = seg.get('speaker', '')
        row[2].text = seg.get('content', '')

    try:
        table.columns[0].width = Inches(0.8)
        table.columns[1].width = Inches(1.2)
        table.columns[2].width = Inches(4.5)
    except Exception:
        # 某些 python-docx 版本不支援直接設定 column.width，忽略即可
        pass

    file_stream = io.BytesIO()
    document.save(file_stream)
    file_stream.seek(0)
    return file_stream


def create_bilingual_transcript_docx(
    original_transcript: list,
    translated_transcript: list,
    meeting_info: dict,
    target_language: str
) -> io.BytesIO:
    """從逐字稿資料創建雙語 Word 文件。"""
    document = Document()
    # 顯示與檔名一律沿用傳入 locale，避免自定義映射
    target_language_name = _safe_language_label(target_language)

    document.add_heading(f"雙語逐字稿：{meeting_info.get('original_filename', 'N/A')}", level=0)
    document.add_paragraph(f"翻譯語言：{target_language_name}")

    table = document.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '時間'
    hdr_cells[1].text = '原文'
    hdr_cells[2].text = f'翻譯 ({target_language_name})'

    translated_map = {i: item.get('translated_content') for i, item in enumerate(translated_transcript)}

    for segment in original_transcript:
        row_cells = table.add_row().cells
        row_cells[0].text = segment.get('start_time_formatted', '')
        
        original_content = f"{segment.get('speaker', '')}: {segment.get('content', '')}"
        row_cells[1].text = original_content
        
        translated_content = translated_map.get(segment.get('id'), '翻譯內容未提供')
        row_cells[2].text = translated_content

    table.columns[0].width = Inches(0.8)
    table.columns[1].width = Inches(3.0)
    table.columns[2].width = Inches(3.0)

    file_stream = io.BytesIO()
    document.save(file_stream)
    file_stream.seek(0)
    return file_stream


def create_bilingual_transcript_txt(
    original_transcript: list,
    translated_transcript: list,
    meeting_info: dict,
    target_language: str
) -> str:
    """從逐字稿資料創建雙語純文字檔案。"""
    target_language_name = _safe_language_label(target_language)
    
    lines = []
    lines.append(f"雙語逐字稿：{meeting_info.get('original_filename', 'N/A')}")
    lines.append(f"翻譯語言：{target_language_name}")
    lines.append("="*40 + "\n")
    
    translated_map = {i: item.get('translated_content') for i, item in enumerate(translated_transcript)}
    
    for segment in original_transcript:
        lines.append(f"[{segment.get('start_time_formatted', '')}] {segment.get('speaker', '')}")
        lines.append(f"  原文: {segment.get('content', '')}")
        lines.append(f"  翻譯: {translated_map.get(segment.get('id'), '翻譯內容未提供')}\n")
        
    return "\n".join(lines) 